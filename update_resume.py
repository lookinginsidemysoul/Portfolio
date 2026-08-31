from docx import Document

SOURCE = "Pallab_Mukherjee_Resume.docx"

replacements = {
    "Senior Business Analyst | BI & Data Visualization | AI-Driven Analytics":
        "Senior Business Analyst | Enterprise BI, AI Systems & Decision Support",
    "Senior Business Analyst with 7+ years of experience across Business Intelligence, data visualization, and applied AI/ML, specializing in Power BI, Tableau, SQL Server, and Python. Proven track record designing scalable data models and interactive dashboards that drive C-suite decision-making, and pioneering AI-augmented analytics tools that reduce manual reporting effort. Recently built a multi-agent SQL assistant that lets business users query data in plain language, saving 3.2 FTEs previously spent on manual SQL reporting. Skilled in leading cross-functional Agile teams, stakeholder engagement, and translating complex datasets into measurable business impact.":
        "Senior Business Analyst with 8+ years of experience across enterprise Business Intelligence, data visualization, and applied AI/ML, specializing in Power BI, Tableau, SQL Server, and Python. Proven track record designing trusted data models and executive dashboards that drive decisions, while building AI-enabled analytics workflows that reduce manual reporting effort. Built a multi-agent SQL assistant that lets business users ask questions in plain language, releasing 3.2 FTEs of reporting capacity. Skilled in leading cross-functional Agile teams, stakeholder engagement, analytics governance, and translating complex datasets into measurable business impact.",
    "BI & Visualization: Power BI, Tableau, Qlik Sense, DAX":
        "BI & Visualization: Power BI, Tableau, Qlik Sense, DAX, Executive Dashboard Design, Data Storytelling",
    "Data & AI: SQL Server, Python, R, Predictive Analytics, Machine Learning, Multi-Agent AI Systems, Prompt Engineering":
        "Data & AI: SQL Server, Python, R, Predictive Analytics, Machine Learning, Text-to-SQL & Multi-Agent Systems, Prompt Design",
    "Delivery: ETL & Data Modeling, Agile/Scrum, Jira, MS Project, Stakeholder & C-Suite Reporting":
        "Data Architecture & Delivery: Dimensional Modeling, ETL, Data Quality, Data Lineage, Agile/Scrum, Jira, MS Project",
    "Leadership: Cross-functional Team Leadership, Mentoring, Process Optimization, Project Documentation":
        "Governance & Leadership: Semantic Models, DAX Optimization, Row-Level Security, Stakeholder & C-Suite Reporting, Team Leadership, Mentoring",
    "Built and deployed a multi-agent SQL assistant enabling business users to run autonomous SQL queries through simple chat input, eliminating manual query-writing and saving 3.2 FTEs in reporting effort.":
        "Built and deployed a multi-agent SQL assistant enabling business users to ask data questions in plain language; the workflow plans and executes SQL behind the scenes, releasing 3.2 FTEs of manual reporting capacity.",
    "Administered Power BI and Tableau platforms; mentored a BI team of four in best practices, code reviews, and performance tuning, resulting in 20+ interactive dashboards that increased stakeholder engagement by 30%.":
        "Administered Power BI and Tableau platforms; mentored a BI team of four in dashboard standards, code reviews, and performance tuning, delivering 20+ interactive dashboards that increased stakeholder engagement by 30%.",
}

def replace_in_paragraph(paragraph):
    text = paragraph.text.strip()
    replacement = replacements.get(text)
    if replacement is None:
        return False
    paragraph.clear()
    paragraph.add_run(replacement)
    return True

document = Document(SOURCE)
changed = 0
for paragraph in document.paragraphs:
    changed += replace_in_paragraph(paragraph)
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                changed += replace_in_paragraph(paragraph)

if changed != len(replacements):
    missing = len(replacements) - changed
    raise RuntimeError(f"Expected {len(replacements)} replacements; completed {changed} (missing {missing}).")

document.save(SOURCE)
print(f"Updated {changed} resume sections.")
